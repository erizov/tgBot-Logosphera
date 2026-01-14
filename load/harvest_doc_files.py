"""
Сбор цитат из локальных doc/docx файлов, начинающихся с "aph".
Парсит только цитаты из 1-5 предложений до слова "Here".
"""

import os
import json
import re
import logging
import glob
from typing import List, Dict, Optional, Tuple
from pathlib import Path

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import textract
    TEXTRACT_AVAILABLE = True
except (ImportError, Exception):
    TEXTRACT_AVAILABLE = False

try:
    import subprocess
    # Проверяем наличие antiword (внешняя утилита)
    result = subprocess.run(
        ['antiword', '-v'],
        capture_output=True,
        timeout=2
    )
    ANTIWORD_AVAILABLE = result.returncode == 0
except (FileNotFoundError, Exception):
    ANTIWORD_AVAILABLE = False

logging.basicConfig(
    format='%(asctime)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)


def is_valid_quotation(text: str) -> bool:
    """
    Проверка валидности цитаты (строгие фильтры).

    Args:
        text: Текст цитаты

    Returns:
        True если цитата валидна
    """
    if not text or len(text.strip()) < 20:
        return False

    # Убираем лишние пробелы
    text = ' '.join(text.split())

    # НЕ допускаем:
    # - Цифры (арабские и римские)
    if re.search(r'\d', text):
        return False

    # - Римские цифры (кроме 'I' как местоимения)
    roman_numerals = re.search(r'\b[IVXLCDM]{2,}\b', text, re.IGNORECASE)
    if roman_numerals:
        return False

    # - Имена людей (два подряд заглавных слова)
    if re.search(r'\b[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+', text):
        # Исключаем начало предложения
        if not text[0].isupper():
            return False

    # - Адреса, места
    place_keywords = [
        r'\bулица\b', r'\bпроспект\b', r'\bпереулок\b',
        r'\bмосква\b', r'\bлондон\b', r'\bпариж\b',
        r'\bstreet\b', r'\bavenue\b', r'\broad\b'
    ]
    for keyword in place_keywords:
        if re.search(keyword, text, re.IGNORECASE):
            return False

    # - Названия книг, фильмов (в кавычках с заглавными буквами)
    if re.search(r'["«»""][А-ЯЁ][а-яё]+.*[А-ЯЁ][а-яё]+["«»""]', text):
        return False

    # - URL и email
    if re.search(r'http[s]?://|www\.|@', text, re.IGNORECASE):
        return False

    # - Даты
    month_names = [
        r'\bянварь\b', r'\bфевраль\b', r'\bмарт\b', r'\bапрель\b',
        r'\bмай\b', r'\bиюнь\b', r'\bиюль\b', r'\bавгуст\b',
        r'\bсентябрь\b', r'\bоктябрь\b', r'\bноябрь\b', r'\bдекабрь\b',
        r'\bjanuary\b', r'\bfebruary\b', r'\bmarch\b', r'\bapril\b'
    ]
    for month in month_names:
        if re.search(month, text, re.IGNORECASE):
            return False

    # - Театральные ссылки
    theater_keywords = [
        r'\bакт\b', r'\bсцена\b', r'\bстраница\b', r'\bглава\b',
        r'\bact\b', r'\bscene\b', r'\bpage\b', r'\bchapter\b'
    ]
    for keyword in theater_keywords:
        if re.search(keyword, text, re.IGNORECASE):
            return False

    # - Спам паттерны (повторяющиеся символы)
    if re.search(r'(.)\1{4,}', text):
        return False

    return True


def clean_text(text: str) -> str:
    """
    Очистка текста от лишних символов.

    Args:
        text: Исходный текст

    Returns:
        Очищенный текст
    """
    if not text:
        return ""

    # Убираем лишние пробелы
    text = ' '.join(text.split())

    # Убираем множественные переносы строк
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Убираем лишние точки и запятые
    text = re.sub(r'\.{3,}', '...', text)
    text = re.sub(r',{2,}', ',', text)

    return text.strip()


def count_sentences(text: str) -> int:
    """
    Подсчет количества предложений в тексте.

    Args:
        text: Текст

    Returns:
        Количество предложений
    """
    # Разделяем по знакам окончания предложения
    sentences = re.split(r'[.!?]+', text)
    # Фильтруем пустые строки
    sentences = [s.strip() for s in sentences if s.strip()]
    return len(sentences)


def extract_author(text: str) -> Tuple[str, Optional[str]]:
    """
    Извлечение автора из текста цитаты.

    Args:
        text: Текст цитаты с возможным автором

    Returns:
        Кортеж (текст цитаты, автор или None)
    """
    # Паттерны для автора: "— Author", "– Author", "—Author", etc.
    author_patterns = [
        r'\s*[—–]\s*([А-ЯЁ][А-ЯЁа-яё\s\.]+?)(?:\.|$|\n)',
        r'\s*[—–]\s*([A-Z][A-Za-z\s\.]+?)(?:\.|$|\n)',
        r'\s*—\s*([А-ЯЁ][А-ЯЁа-яё\s\.]+?)(?:\.|$|\n)',
        r'\s*—\s*([A-Z][A-Za-z\s\.]+?)(?:\.|$|\n)',
    ]

    author = None
    quote_text = text

    for pattern in author_patterns:
        match = re.search(pattern, text)
        if match:
            author = match.group(1).strip()
            # Убираем автора из текста цитаты
            quote_text = text[:match.start()].strip()
            # Проверяем, что автор не слишком длинный (не более 100 символов)
            if len(author) > 100:
                author = None
                quote_text = text
            else:
                break

    return quote_text, author


def read_docx_file(file_path: str) -> str:
    """
    Чтение текста из docx файла.

    Args:
        file_path: Путь к файлу

    Returns:
        Текст файла
    """
    if not DOCX_AVAILABLE:
        logger.warning(
            f"python-docx not available, skipping {file_path}"
        )
        return ""

    try:
        doc = Document(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        return '\n'.join(paragraphs)
    except Exception as e:
        logger.error(f"Error reading docx file {file_path}: {e}")
        return ""


def read_doc_file(file_path: str) -> str:
    """
    Чтение текста из doc файла (старый формат).

    Args:
        file_path: Путь к файлу

    Returns:
        Текст файла
    """
    # Сначала пробуем python-docx (на случай если файл на самом деле .docx)
    if DOCX_AVAILABLE:
        try:
            doc = Document(file_path)
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            if paragraphs:
                logger.info(
                    f"Successfully read {file_path} as docx format "
                    "(file might be misnamed)"
                )
                return '\n'.join(paragraphs)
        except Exception:
            # Не docx формат, продолжаем
            pass

    # Проверяем, является ли файл бинарным OLE2 форматом
    try:
        with open(file_path, 'rb') as f:
            header = f.read(8)
            # OLE2 формат начинается с D0CF11E0A1B11AE1
            if header[:8] == b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1':
                # Пробуем использовать antiword если доступен
                if ANTIWORD_AVAILABLE:
                    try:
                        logger.info(
                            f"Attempting to read binary .doc file "
                            f"{file_path} using antiword"
                        )
                        result = subprocess.run(
                            ['antiword', file_path],
                            capture_output=True,
                            text=True,
                            timeout=30,
                            encoding='utf-8',
                            errors='replace'
                        )
                        if result.returncode == 0 and result.stdout:
                            text = result.stdout.strip()
                            if text and len(text) > 10:
                                logger.info(
                                    f"Successfully extracted text from "
                                    f"{file_path} using antiword"
                                )
                                return text
                    except Exception as e:
                        logger.warning(
                            f"antiword failed to read {file_path}: {e}"
                        )
                
                # Пробуем использовать textract если доступен
                if TEXTRACT_AVAILABLE:
                    try:
                        logger.info(
                            f"Attempting to read binary .doc file "
                            f"{file_path} using textract"
                        )
                        text = textract.process(file_path).decode('utf-8')
                        if text and len(text.strip()) > 10:
                            logger.info(
                                f"Successfully extracted text from "
                                f"{file_path} using textract"
                            )
                            return text
                    except Exception as e:
                        logger.warning(
                            f"textract failed to read {file_path}: {e}"
                        )
                
                logger.warning(
                    f"{file_path} is a binary OLE2 .doc file. "
                    "Old .doc format requires additional tools. "
                    "Options:\n"
                    "  1. Convert .doc to .docx using Microsoft Word or LibreOffice\n"
                    "  2. Install antiword (external utility): "
                    "https://www.winfield.demon.nl/\n"
                    "  3. Use textract (may have compatibility issues with pip>=24.1): "
                    "pip install textract (requires pip<24.1)"
                )
                return ""
    except Exception:
        pass

    # Пробуем разные кодировки (на случай если это текстовый файл)
    encodings = ['utf-8', 'cp1251', 'windows-1251', 'latin-1', 'cp866']

    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                content = f.read()
                # Проверяем, что получили читаемый текст
                # Проверяем наличие русских букв или обычного текста
                if content and len(content.strip()) > 10:
                    # Проверяем, что это не просто бинарные данные
                    readable_chars = sum(
                        1 for c in content[:1000]
                        if c.isprintable() or c.isspace()
                    )
                    if readable_chars > len(content[:1000]) * 0.7:
                        return content
        except UnicodeDecodeError:
            continue
        except Exception as e:
            logger.debug(
                f"Error reading doc file {file_path} with "
                f"encoding {encoding}: {e}"
            )
            continue

    # Если не получилось прочитать как текст, пробуем бинарный режим
    try:
        with open(file_path, 'rb') as f:
            content = f.read()
            # Пробуем декодировать как текст
            for encoding in encodings:
                try:
                    text = content.decode(encoding, errors='ignore')
                    if text and len(text.strip()) > 10:
                        # Проверяем читаемость
                        readable_chars = sum(
                            1 for c in text[:1000]
                            if c.isprintable() or c.isspace()
                        )
                        if readable_chars > len(text[:1000]) * 0.7:
                            return text
                except:
                    continue
    except Exception as e:
        logger.error(f"Error reading doc file {file_path}: {e}")

    logger.warning(
        f"Could not read doc file {file_path}. "
        "File might be in binary OLE2 format (old .doc). "
        "Consider converting to .docx or installing textract library."
    )
    return ""


def parse_text_up_to_here(text: str) -> str:
    """
    Извлечение текста до слова "Here" (регистронезависимо).

    Args:
        text: Исходный текст

    Returns:
        Текст до слова "Here"
    """
    # Ищем слово "Here" (регистронезависимо)
    # Учитываем возможные варианты: "Here", "here", "HERE"
    match = re.search(
        r'\b[Hh][Ee][Rr][Ee]\b',
        text,
        re.IGNORECASE
    )

    if match:
        return text[:match.start()].strip()

    return text.strip()


def extract_quotes_from_text(text: str) -> List[Dict]:
    """
    Извлечение цитат из текста.

    Args:
        text: Исходный текст

    Returns:
        Список цитат
    """
    quotes = []

    # Извлекаем текст до слова "Here"
    text = parse_text_up_to_here(text)

    if not text:
        return quotes

    # Разделяем текст на предложения
    # Используем более сложный паттерн для разделения предложений
    sentences = re.split(
        r'(?<=[.!?])\s+(?=[А-ЯЁA-Z])',
        text
    )

    # Группируем предложения в цитаты (1-5 предложений)
    current_quote = []
    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue

        # Если текущая цитата уже содержит 5 предложений, сохраняем её
        if len(current_quote) >= 5:
            quote_text = ' '.join(current_quote)
            quote_text = clean_text(quote_text)

            # Извлекаем автора
            quote_text, author = extract_author(quote_text)

            # Проверяем валидность
            if is_valid_quotation(quote_text):
                # Проверяем, что это русский текст
                if re.search(r'[а-яёА-ЯЁ]', quote_text):
                    quotes.append({
                        "text": quote_text,
                        "author": author,
                        "source": "local_doc_files",
                        "lang": "ru"
                    })

            current_quote = [sentence]
        else:
            current_quote.append(sentence)

            # Если цитата содержит от 1 до 5 предложений, проверяем её
            if 1 <= len(current_quote) <= 5:
                quote_text = ' '.join(current_quote)
                quote_text = clean_text(quote_text)

                # Извлекаем автора
                quote_text, author = extract_author(quote_text)

                # Проверяем валидность
                if is_valid_quotation(quote_text):
                    # Проверяем, что это русский текст
                    if re.search(r'[а-яёА-ЯЁ]', quote_text):
                        quotes.append({
                            "text": quote_text,
                            "author": author,
                            "source": "local_doc_files",
                            "lang": "ru"
                        })

    # Обрабатываем последнюю цитату, если она не была обработана
    if current_quote and len(current_quote) <= 5:
        quote_text = ' '.join(current_quote)
        quote_text = clean_text(quote_text)

        # Извлекаем автора
        quote_text, author = extract_author(quote_text)

        # Проверяем валидность
        if is_valid_quotation(quote_text):
            # Проверяем, что это русский текст
            if re.search(r'[а-яёА-ЯЁ]', quote_text):
                quotes.append({
                    "text": quote_text,
                    "author": author,
                    "source": "local_doc_files",
                    "lang": "ru"
                })

    return quotes


def harvest_doc_files(
    folder_path: str = "data",
    output_file: str = "data/doc_files.json"
) -> List[Dict]:
    """
    Сбор цитат из doc/docx файлов, начинающихся с "aph".

    Args:
        folder_path: Путь к папке с файлами
        output_file: Имя файла для сохранения

    Returns:
        Список цитат
    """
    quotes = []
    folder = Path(folder_path)

    # Ищем все файлы, начинающиеся с "aph"
    docx_files = list(folder.glob("aph*.docx"))
    doc_files = list(folder.glob("aph*.doc"))

    all_files = docx_files + doc_files

    if not all_files:
        logger.warning(
            f"No files starting with 'aph' found in {folder_path}"
        )
        return quotes

    logger.info(f"Found {len(all_files)} files starting with 'aph'")

    for file_path in all_files:
        logger.info(f"Processing {file_path.name}...")

        try:
            # Читаем файл в зависимости от расширения
            if file_path.suffix.lower() == '.docx':
                text = read_docx_file(str(file_path))
            elif file_path.suffix.lower() == '.doc':
                text = read_doc_file(str(file_path))
            else:
                logger.warning(
                    f"Unknown file extension for {file_path.name}"
                )
                continue

            if not text:
                logger.warning(f"No text extracted from {file_path.name}")
                continue

            # Извлекаем цитаты
            file_quotes = extract_quotes_from_text(text)

            if file_quotes:
                quotes.extend(file_quotes)
                logger.info(
                    f"Extracted {len(file_quotes)} quotes from "
                    f"{file_path.name}"
                )
            else:
                logger.warning(
                    f"No valid quotes found in {file_path.name}"
                )

        except Exception as e:
            logger.error(f"Error processing {file_path.name}: {e}")
            continue

    # Дедупликация по тексту
    seen_texts = set()
    unique_quotes = []
    for quote in quotes:
        text_lower = quote["text"].strip().lower()
        if text_lower not in seen_texts:
            seen_texts.add(text_lower)
            unique_quotes.append(quote)

    logger.info(
        f"Total quotes extracted: {len(quotes)}, "
        f"unique: {len(unique_quotes)}"
    )

    # Сохраняем в файл
    if unique_quotes:
        with open(output_file, "w", encoding="utf-8") as f:
            json.dump(unique_quotes, f, ensure_ascii=False, indent=2)
        logger.info(f"Saved {len(unique_quotes)} quotes to {output_file}")
    else:
        # Сохраняем пустой список
        with open(output_file, "w", encoding="utf-8") as f:
            json.dump([], f, ensure_ascii=False, indent=2)
        logger.info(f"No quotes to save, created empty {output_file}")

    return unique_quotes


if __name__ == "__main__":
    harvest_doc_files()
